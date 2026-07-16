"""Document parser service - PyMuPDF for PDF, python-docx for DOCX, plain read for TXT/MD"""
import base64
import logging
import time
from pathlib import Path
from dataclasses import dataclass

import fitz  # PyMuPDF

logger = logging.getLogger(__name__)

# Minimum image dimension in pixels - skip tiny icons/logos
_MIN_IMAGE_DIM = 50


@dataclass
class ParsedDocument:
    """Parsed document result"""
    markdown: str  # Markdown content
    images: list[dict]  # [{"caption": str, "base64": str}]
    tables: list[dict]  # [{"caption": str, "markdown": str}]


class DocumentParser:
    """Document parser - PyMuPDF for PDF, python-docx for DOCX, plain read for TXT/MD"""

    def parse(self, file_path: str, extract_images: bool = True) -> ParsedDocument:
        """
        Parse document, extract markdown content and multimodal elements.
        Supports PDF, DOCX, TXT, MD formats.

        Args:
            file_path: Document file path
            extract_images: Whether to extract images from PDF (default True)
        """
        suffix = Path(file_path).suffix.lower()

        if suffix == ".pdf":
            return self._parse_pdf_with_pymupdf(file_path, extract_images)
        elif suffix in (".docx", ".doc"):
            return self._parse_docx(file_path)
        elif suffix in (".txt", ".md"):
            return self._parse_plaintext(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def _parse_pdf_with_pymupdf(self, file_path: str, extract_images: bool) -> ParsedDocument:
        """
        Parse PDF using PyMuPDF - read embedded text directly (no OCR),
        extract images and tables per page.

        Key: detect valid tables first, exclude their bbox from text extraction
        to avoid duplicate content. Layout-only tables (garbage) are ignored so
        their text stays in the normal text flow.
        """
        logger.info(f"Parsing PDF with PyMuPDF: {file_path} | extract_images={extract_images}")
        start_time = time.monotonic()

        doc = fitz.open(file_path)
        text_parts: list[str] = []
        images: list[dict] = []
        tables: list[dict] = []

        for page_num in range(len(doc)):
            page = doc[page_num]

            # 1. Find valid tables and collect their bboxes for text exclusion
            page_tables, valid_table_rects = self._extract_page_tables(page, page_num)
            tables.extend(page_tables)

            # 2. Extract text, excluding areas covered by valid tables
            page_text = self._extract_text_excluding_rects(page, valid_table_rects)
            if page_text.strip():
                text_parts.append(page_text.strip())

            # 3. Extract images (skip tiny icons)
            if extract_images:
                page_images = self._extract_page_images(doc, page, page_num)
                images.extend(page_images)

        doc.close()

        # Build final markdown: text content + valid tables inline
        markdown_content = "\n\n".join(text_parts) if text_parts else ""

        # Append valid table markdown at the end
        if tables:
            table_section = "\n\n---\n"
            for t in tables:
                table_section += f"\n### {t['caption']}\n\n{t['markdown']}\n"
            markdown_content = markdown_content + table_section

        total_elapsed = time.monotonic() - start_time
        logger.info(
            f"PDF parsed (PyMuPDF): "
            f"markdown={len(markdown_content)} chars | "
            f"images={len(images)} | tables={len(tables)} | "
            f"total={total_elapsed:.2f}s"
        )

        return ParsedDocument(
            markdown=markdown_content,
            images=images,
            tables=tables,
        )

    @staticmethod
    def _extract_page_images(
        doc: fitz.Document, page: fitz.Page, page_num: int,
    ) -> list[dict]:
        """Extract meaningful images from a PDF page, skip tiny icons."""
        images: list[dict] = []

        for img_info in page.get_images(full=True):
            xref = img_info[0]
            try:
                img_data = doc.extract_image(xref)
                if not img_data:
                    continue

                width = img_data.get("width", 0)
                height = img_data.get("height", 0)

                # Skip tiny icons/logos
                if width < _MIN_IMAGE_DIM or height < _MIN_IMAGE_DIM:
                    continue

                img_bytes = img_data.get("image")
                if not img_bytes:
                    continue

                b64 = base64.b64encode(img_bytes).decode("utf-8")
                ext = img_data.get("ext", "png")
                images.append({
                    "caption": f"Page {page_num + 1} Image ({width}x{height})",
                    "base64": b64,
                    "mime_type": f"image/{ext}",
                })
            except Exception as e:
                logger.debug(f"Failed to extract image xref={xref} on page {page_num + 1}: {e}")
                continue

        return images

    @staticmethod
    def _extract_page_tables(
        page: fitz.Page, page_num: int,
    ) -> tuple[list[dict], list[fitz.Rect]]:
        """
        Extract valid data tables from a PDF page.
        Returns (tables, valid_rects) - valid_rects are bboxes of accepted tables,
        used to exclude their text from normal text extraction.

        Filter rules:
        - Skip tables where majority of column names are auto-generated (Col0, Col1...)
        - Skip tables with fewer than 2 data rows
        - Skip tables with only 1 column (just a text block)
        """
        tables: list[dict] = []
        valid_rects: list[fitz.Rect] = []

        try:
            found_tables = page.find_tables()
            for idx, tab in enumerate(found_tables):
                try:
                    df = tab.to_pandas()
                    if df.empty:
                        continue

                    cols = list(df.columns)
                    num_cols = len(cols)

                    # Skip single-column "tables" - just text blocks
                    if num_cols < 2:
                        continue

                    # Skip tables where most columns are auto-named (Col0, Col1...)
                    auto_col_count = sum(1 for c in cols if str(c).startswith("Col"))
                    if auto_col_count > num_cols / 2:
                        continue

                    # Skip tables with fewer than 2 data rows
                    if len(df) < 2:
                        continue

                    md = df.to_markdown(index=False)
                    tables.append({
                        "caption": f"Page {page_num + 1} Table {idx + 1}",
                        "markdown": md,
                    })
                    valid_rects.append(fitz.Rect(tab.bbox))

                except Exception as e:
                    logger.debug(f"Failed to convert table {idx} on page {page_num + 1}: {e}")
                    continue
        except Exception as e:
            logger.debug(f"Table detection failed on page {page_num + 1}: {e}")

        return tables, valid_rects

    @staticmethod
    def _extract_text_excluding_rects(
        page: fitz.Page, exclude_rects: list[fitz.Rect],
    ) -> str:
        """
        Extract text from page, skipping text blocks that fall inside
        any of the exclude_rects (valid table areas).
        """
        if not exclude_rects:
            return page.get_text("text")

        # Use text blocks to filter by position
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        text_parts: list[str] = []

        for block in blocks:
            if block["type"] != 0:  # only text blocks
                continue

            block_rect = fitz.Rect(block["bbox"])

            # Skip blocks that overlap with any valid table rect
            in_table = any(
                block_rect.intersects(tr) for tr in exclude_rects
            )
            if in_table:
                continue

            # Collect text from this block's lines
            for line in block["lines"]:
                line_text = "".join(span["text"] for span in line["spans"])
                if line_text.strip():
                    text_parts.append(line_text.strip())

        return "\n".join(text_parts)

    @staticmethod
    def _parse_docx(file_path: str) -> ParsedDocument:
        """Parse DOCX using python-docx - lightweight, no OCR dependency."""
        from docx import Document as DocxDocument

        logger.info(f"Parsing DOCX with python-docx: {file_path}")
        start_time = time.monotonic()

        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        markdown_content = "\n\n".join(paragraphs)

        total_elapsed = time.monotonic() - start_time
        logger.info(
            f"DOCX parsed: markdown={len(markdown_content)} chars | "
            f"total={total_elapsed:.2f}s"
        )

        return ParsedDocument(
            markdown=markdown_content,
            images=[],
            tables=[],
        )

    @staticmethod
    def _parse_plaintext(file_path: str) -> ParsedDocument:
        """Parse TXT/MD files - just read the content directly."""
        logger.info(f"Reading plaintext file: {file_path}")

        content = Path(file_path).read_text(encoding="utf-8")

        return ParsedDocument(
            markdown=content,
            images=[],
            tables=[],
        )
